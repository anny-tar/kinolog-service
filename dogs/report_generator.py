# Этот модуль отвечает за генерацию отчётов в формате .docx
# Используем библиотеку python-docx

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import date
import io  # для работы с файлами в памяти (без сохранения на диск)


def set_cell_style(cell, text, bold=False, center=False, bg_color=None):
    """
    Вспомогательная функция — применяет стиль к ячейке таблицы.
    Вынесена отдельно чтобы не повторять код в каждой таблице.
    """
    cell.text = text
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Закрашиваем фон ячейки если передан цвет (шапка таблицы)
    if bg_color:
        from docx.oxml.ns import qn
        from docx.oxml import parse_xml
        shading = parse_xml(
            f'<w:shd {chr(10)}xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:val="clear" w:color="auto" w:fill="{bg_color}"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)


def setup_document():
    """
    Создаёт базовый документ с настройками страницы.
    Возвращает готовый объект Document.
    """
    doc = Document()

    # Настраиваем поля страницы (в сантиметрах)
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

    # Устанавливаем шрифт по умолчанию для всего документа
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    return doc


def add_header(doc, title, date_from=None, date_to=None):
    """Добавляет шапку отчёта: название, период, дата формирования"""

    # Название организации
    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org.add_run('ГУ МВД России по Самарской области')
    run.bold = True
    run.font.size = Pt(12)

    dept = doc.add_paragraph()
    dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = dept.add_run('Центр кинологической службы')
    run2.font.size = Pt(11)

    # Разделитель
    doc.add_paragraph()

    # Название отчёта
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = heading.add_run(title.upper())
    run3.bold = True
    run3.font.size = Pt(14)

    # Период отчёта
    if date_from and date_to:
        period = doc.add_paragraph()
        period.alignment = WD_ALIGN_PARAGRAPH.CENTER
        period.add_run(
            f'за период с {date_from.strftime("%d.%m.%Y")} по {date_to.strftime("%d.%m.%Y")}'
        ).font.size = Pt(11)

    # Дата формирования (справа)
    date_par = doc.add_paragraph()
    date_par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_par.add_run(f'Дата формирования: {date.today().strftime("%d.%m.%Y")}').font.size = Pt(10)

    doc.add_paragraph()  # пустая строка


# ==============================================================================
# ОТЧЁТ 1: Список служебных собак
# ==============================================================================

def generate_dogs_report(date_from=None, date_to=None):
    """
    Генерирует отчёт по всем служебным собакам.
    Возвращает файл в памяти (BytesIO) — не сохраняет на диск.
    """
    from .models import ServiceDog

    doc = setup_document()
    add_header(doc, 'Отчёт по служебным собакам', date_from, date_to)

    # Получаем данных из БД
    dogs = ServiceDog.objects.select_related('status', 'main_kennel').order_by('name')

    # Итоговая статистика перед таблицей
    total = dogs.count()
    active = dogs.filter(status__name='В работе').count()

    summary = doc.add_paragraph()
    summary.add_run(f'Всего собак: {total}. Из них в работе: {active}.').font.size = Pt(11)
    doc.add_paragraph()

    # Создаём таблицу: количество строк = собаки + 1 (шапка)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Шапка таблицы
    headers = ['№', 'Кличка', 'Инв. номер', 'Порода', 'Статус', 'Кинолог']
    for i, header in enumerate(headers):
        set_cell_style(table.rows[0].cells[i], header, bold=True, center=True, bg_color='D9E1F2')

    # Данные
    for idx, dog in enumerate(dogs, 1):
        row = table.add_row().cells
        set_cell_style(row[0], str(idx), center=True)
        set_cell_style(row[1], dog.name)
        set_cell_style(row[2], dog.inventory_number)
        set_cell_style(row[3], dog.breed)
        set_cell_style(row[4], str(dog.status))
        set_cell_style(row[5], str(dog.main_kennel) if dog.main_kennel else '—')

    # Подпись руководителя в конце
    _add_signature(doc)

    # Сохраняем документ в памяти и возвращаем
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==============================================================================
# ОТЧЁТ 2: Журнал тренировок за период
# ==============================================================================

def generate_trainings_report(date_from, date_to):
    """Отчёт по тренировкам за выбранный период"""
    from .models import Training

    doc = setup_document()
    add_header(doc, 'Журнал тренировочных занятий', date_from, date_to)

    # Фильтруем тренировки по периоду
    trainings = Training.objects.filter(
        datetime__date__gte=date_from,
        datetime__date__lte=date_to
    ).select_related('dog', 'kennel', 'skill').order_by('datetime')

    doc.add_paragraph(f'Всего занятий за период: {trainings.count()}')
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    headers = ['№', 'Дата', 'Кличка собаки', 'Кинолог', 'Навык', 'Оценка']
    for i, h in enumerate(headers):
        set_cell_style(table.rows[0].cells[i], h, bold=True, center=True, bg_color='D9E1F2')

    for idx, t in enumerate(trainings, 1):
        row = table.add_row().cells
        set_cell_style(row[0], str(idx), center=True)
        set_cell_style(row[1], t.datetime.strftime('%d.%m.%Y'))
        set_cell_style(row[2], t.dog.name)
        set_cell_style(row[3], t.kennel.full_name if t.kennel else '—')
        set_cell_style(row[4], t.skill.name)
        set_cell_style(row[5], str(t.score) if t.score else '—', center=True)

    _add_signature(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==============================================================================
# ОТЧЁТ 3: Ветеринарные мероприятия за период
# ==============================================================================

def generate_vet_report(date_from, date_to):
    """Отчёт по ветеринарным процедурам за период"""
    from .models import VeterinaryRecord

    doc = setup_document()
    add_header(doc, 'Отчёт по ветеринарным мероприятиям', date_from, date_to)

    records = VeterinaryRecord.objects.filter(
        procedure_date__gte=date_from,
        procedure_date__lte=date_to
    ).select_related('dog', 'procedure_type', 'veterinarian').order_by('procedure_date')

    routine = records.filter(is_routine=True).count()
    urgent = records.filter(is_routine=False).count()

    doc.add_paragraph(f'Всего процедур: {records.count()}. Плановых: {routine}. Внеплановых: {urgent}.')
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    headers = ['№', 'Дата', 'Собака', 'Процедура', 'Ветеринар', 'Плановая']
    for i, h in enumerate(headers):
        set_cell_style(table.rows[0].cells[i], h, bold=True, center=True, bg_color='D9E1F2')

    for idx, r in enumerate(records, 1):
        row = table.add_row().cells
        set_cell_style(row[0], str(idx), center=True)
        set_cell_style(row[1], r.procedure_date.strftime('%d.%m.%Y'))
        set_cell_style(row[2], r.dog.name)
        set_cell_style(row[3], r.procedure_type.name)
        set_cell_style(row[4], r.veterinarian.full_name if r.veterinarian else '—')
        set_cell_style(row[5], 'Да' if r.is_routine else 'Нет', center=True)

    _add_signature(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==============================================================================
# ОТЧЁТ 4: Служебные мероприятия за период
# ==============================================================================

def generate_events_report(date_from, date_to):
    """Отчёт по служебным мероприятиям за период"""
    from .models import ServiceEvent

    doc = setup_document()
    add_header(doc, 'Отчёт по служебным мероприятиям', date_from, date_to)

    events = ServiceEvent.objects.filter(
        datetime__date__gte=date_from,
        datetime__date__lte=date_to
    ).select_related('dog', 'kennel', 'event_type').order_by('datetime')

    doc.add_paragraph(f'Всего мероприятий за период: {events.count()}.')
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    headers = ['№', 'Дата', 'Тип', 'Собака / Кинолог', 'Место']
    for i, h in enumerate(headers):
        set_cell_style(table.rows[0].cells[i], h, bold=True, center=True, bg_color='D9E1F2')

    for idx, e in enumerate(events, 1):
        row = table.add_row().cells
        set_cell_style(row[0], str(idx), center=True)
        set_cell_style(row[1], e.datetime.strftime('%d.%m.%Y'))
        set_cell_style(row[2], e.event_type.name)
        kennel_name = e.kennel.full_name if e.kennel else '—'
        set_cell_style(row[3], f'{e.dog.name} / {kennel_name}')
        set_cell_style(row[4], e.location)

    _add_signature(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==============================================================================
# ВСПОМОГАТЕЛЬНАЯ ФУНКЦИЯ — блок подписи руководителя
# ==============================================================================

def _add_signature(doc):
    """Добавляет блок подписи в конец документа"""
    doc.add_paragraph()
    doc.add_paragraph()

    sig = doc.add_paragraph()
    sig.add_run('Начальник Центра кинологической службы').font.size = Pt(11)

    line = doc.add_paragraph()
    run = line.add_run('________________________________  /________________/')
    run.font.size = Pt(11)

    hint = doc.add_paragraph()
    hint.add_run('           (подпись)                                                                     (ФИО)').font.size = Pt(9)